import streamlit as st
import pdfplumber
import pandas as pd
from docx import Document
import io
import os
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import re

st.set_page_config(page_title="Ultimate PDF Data Extractor", layout="wide")
st.title("Ultimate PDF Data Extractor")

# Ensure feedback directory exists
os.makedirs("feedback", exist_ok=True)

uploaded_file = st.file_uploader("Upload a PDF", type="pdf")

def clean_text(text):
    # Remove all characters not allowed in XML/Word (tab, CR, LF, BMP unicode)
    def safe_char(c):
        o = ord(c)
        return (
            o == 0x9 or o == 0xA or o == 0xD or
            (0x20 <= o <= 0xD7FF) or (0xE000 <= o <= 0xFFFD)
        )
    text = ''.join(c if safe_char(c) else ' ' for c in text)
    text = text.encode("utf-8", "ignore").decode("utf-8", "ignore")
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

if uploaded_file:
    st.success("PDF uploaded! Processing...")

    # --- READ PDF INTO MEMORY BUFFER ---
    pdf_bytes = uploaded_file.read()

    # --- TEXT & TABLES EXTRACTION (pdfplumber with OCR fallback) ---
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        all_text = ""
        tables = []
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if not text:
                # OCR fallback if no text is found on the page
                images_from_pdf = convert_from_bytes(pdf_bytes, first_page=page_num, last_page=page_num)
                ocr_text = ""
                for img in images_from_pdf:
                    ocr_text += pytesseract.image_to_string(img)
                all_text += f"\n--- Page {page_num} (OCR) ---\n{ocr_text}"
            else:
                all_text += f"\n--- Page {page_num} ---\n{text}"
            page_tables = page.extract_tables() or []
            tables.extend(page_tables)

    # --- IMAGES EXTRACTION (PyMuPDF) ---
    st.subheader("Extracted Images")
    images = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page_num in range(len(doc)):
        for img_index, img in enumerate(doc.get_page_images(page_num)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            images.append(image_bytes)
            st.image(image_bytes, caption=f"Page {page_num+1} Image {img_index+1}")
            st.download_button(
                label=f"Download Image {page_num+1}-{img_index+1}",
                data=image_bytes,
                file_name=f"pdf_image_{page_num+1}_{img_index+1}.png",
                mime="image/png"
            )
    if not images:
        st.info("No images found in this PDF.")

    # --- TEXT SECTION ---
    st.subheader("Extracted Text (editable)")
    text_area = st.text_area("Edit extracted text:", all_text, height=300)

    # --- TABLES SECTION ---
    st.subheader("Extracted Tables")
    edited_tables = []
    if tables:
        for i, table in enumerate(tables):
            try:
                df = pd.DataFrame(table[1:], columns=table[0])
            except Exception:
                df = pd.DataFrame(table)
            edited_df = st.data_editor(df, num_rows="dynamic", key=f"table_{i}")
            edited_tables.append(edited_df)
            st.write("---")
    else:
        st.info("No tables found in this PDF.")

    # --- EXPORT SECTION ---
    st.subheader("Export Extracted Data")

    # Export to Word (.docx)
    doc_buf = io.BytesIO()
    if st.button("Export to Word (.docx)"):
        doc = Document()
        doc.add_heading("Extracted PDF Text", level=1)
        doc.add_paragraph(clean_text(text_area))
        if edited_tables:
            for idx, df in enumerate(edited_tables):
                doc.add_heading(f"Table {idx+1}", level=2)
                t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
                for j, col in enumerate(df.columns):
                    t.cell(0, j).text = clean_text(str(col))
                for i, row in df.iterrows():
                    for j, val in enumerate(row):
                        t.cell(i+1, j).text = clean_text(str(val))
        doc.save(doc_buf)
        st.download_button(
            "Download Word File", doc_buf.getvalue(), "output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Export to Excel (.xlsx)
    excel_buf = io.BytesIO()
    if st.button("Export to Excel (.xlsx)"):
        with pd.ExcelWriter(excel_buf, engine="openpyxl") as writer:
            for idx, df in enumerate(edited_tables):
                sheet_name = f"Table_{idx+1}"
                # Excel is more tolerant, but let's clean just in case
                clean_df = df.applymap(lambda x: clean_text(str(x)))
                clean_df.to_excel(writer, index=False, sheet_name=sheet_name)
        st.download_button(
            "Download Excel File", excel_buf.getvalue(), "output.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    # Export text as .txt
    if st.button("Export Text (.txt)"):
        st.download_button("Download Text File", clean_text(text_area).encode(), "output.txt")

    # --- USER FEEDBACK ---
    st.subheader("Submit Feedback / Corrections")
    feedback = st.text_area("How can we improve this extraction?", key="feedback")
    if st.button("Submit Feedback"):
        with open("feedback/feedback.txt", "a", encoding="utf-8") as f:
            f.write(f"--- Feedback for file: {uploaded_file.name} ---\n{clean_text(text_area)}\n{clean_text(feedback)}\n\n")
        st.success("Thank you for your feedback! This will help improve future extractions.")

else:
    st.info("Please upload a PDF to get started.")

st.markdown("---")
st.caption("v0.3 | Code Generator GPT · With OCR fallback for scanned/image PDFs and bulletproof text cleaning for export.")
